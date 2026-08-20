import re
from pathlib import Path


class FileParseError(Exception):
    """Raised when resume file cannot be parsed."""
    pass


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def normalize_text(text: str) -> str:
    """Collapse blank lines and edge whitespace for stable resume text."""
    lines = [line.strip() for line in (text or "").splitlines() if line.strip()]
    return "\n".join(lines)


_BULLET_RE = re.compile(r"^[\s]*[•▪●○◦·∙‣◦◆◇■□★☆※]\s*")
_SECTION_HEADING_LINE_RE = re.compile(
    r"^(?:#+\s*)?(?:个人(?:信息|简介)|教育(?:背景|经历)|工作经历|"
    r"项目(?:经历|经验)|专业技能|技能清单|证书|荣誉|自我评价|"
    r"summary|education|work\s+experience|projects?|skills|"
    r"certifications?|awards?)[\s:：]*$",
    re.IGNORECASE,
)
_INVISIBLE_RE = re.compile(r"[\u200b-\u200d\ufeff]")
_MULTISPACE_RE = re.compile(r"[ \t]+")
_BLANK_RUN_RE = re.compile(r"\n{3,}")


def clean_resume_markdown(text: str) -> str:
    """Normalize messy PDF/DOCX extraction into stable Markdown-like lines.

    This is intentionally conservative and deterministic: it does not rewrite
    prose or invent section names. It only repairs common extraction noise so
    the first A4 render already looks like an edited resume.
    """
    normalized: list[str] = []
    for raw_line in (text or "").splitlines():
        line = _INVISIBLE_RE.sub("", raw_line).replace("\u3000", " ").strip()
        if not line:
            continue
        line = _MULTISPACE_RE.sub(" ", line)
        if not _SECTION_HEADING_LINE_RE.match(line):
            line = _BULLET_RE.sub("- ", line)
        normalized.append(line)

    cleaned: list[str] = []
    for line in normalized:
        if not cleaned:
            cleaned.append(line)
            continue
        prev = cleaned[-1]
        if _SECTION_HEADING_LINE_RE.match(line) or line.startswith(("-", "#")):
            cleaned.append(line)
            continue
        if prev and not _SECTION_HEADING_LINE_RE.match(prev) and not prev.startswith(("-", "#")):
            if not re.search(r"[。！？!?；;：:，,、.\-—]\s*$", prev):
                if len(line) < 40 and not re.search(r"\d{4}[-/年]\d{1,2}", line):
                    cleaned[-1] = f"{prev} {line}"
                    continue
        cleaned.append(line)

    return _BLANK_RUN_RE.sub("\n\n", "\n".join(cleaned)).strip()


def extract_text(path: Path) -> str:
    if not path.exists():
        raise FileParseError(f"File not found: {path}")

    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise FileParseError(
            f"Unsupported format '{ext}'. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    if ext == ".pdf":
        import fitz
        doc = fitz.open(path)
        pages = []
        for page in doc:
            text = page.get_text()
            if text.strip():
                pages.append(normalize_text(text))
        doc.close()
        return clean_resume_markdown("\n\n".join(pages))

    if ext == ".docx":
        from docx import Document
        doc = Document(str(path))
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return clean_resume_markdown("\n".join(lines))

    # .txt
    try:
        return clean_resume_markdown(path.read_text(encoding="utf-8"))
    except UnicodeDecodeError:
        try:
            return clean_resume_markdown(path.read_text(encoding="gb18030"))
        except UnicodeDecodeError as exc:
            raise FileParseError(
                "Text file encoding is not supported; "
                "please save it as UTF-8 or GB18030"
            ) from exc


_SECTION_HEADING_RE = re.compile(
    r"^(?:#\s*)?(?:个人(?:信息|简介)|教育(?:背景|经历)|工作经历|"
    r"项目(?:经历|经验)|专业技能|技能清单|证书|荣誉|自我评价|"
    r"summary|education|work\s+experience|project|skills|"
    r"certifications?|awards?)$",
    re.IGNORECASE,
)


def structured_resume_sections(text: str) -> dict[str, str]:
    """Group normalized resume text into familiar section buckets."""
    sections: dict[str, list[str]] = {}
    current: str | None = None
    for line in (text or "").splitlines():
        heading = line.strip().lstrip("#").strip()
        if _SECTION_HEADING_RE.match(heading):
            current = heading
            sections.setdefault(current, [])
            continue
        if current is not None and line.strip():
            sections[current].append(line.strip())
    return {
        key: "\n".join(values)
        for key, values in sections.items()
        if values
    }
