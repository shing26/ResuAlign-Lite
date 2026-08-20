from pathlib import Path

import pytest

from resualign.parser import (
    FileParseError,
    clean_resume_markdown,
    extract_text,
    normalize_text,
    structured_resume_sections,
)

HERE = Path(__file__).parent
FIXTURES = HERE / "fixtures"


def test_txt_parse():
    text = extract_text(FIXTURES / "sample.txt")
    assert "Python" in text
    assert "Java" in text
    assert len(text) > 50


def test_txt_gb18030_fallback(tmp_path):
    path = tmp_path / "gbk.txt"
    path.write_bytes("简历：Python 开发".encode("gbk"))
    text = extract_text(path)
    assert "Python" in text
    assert "简历" in text


def test_txt_unsupported_encoding_raises(tmp_path):
    path = tmp_path / "binary.txt"
    path.write_bytes(b"\xff\xfe\xfa\x01")
    with pytest.raises(FileParseError, match="encoding"):
        extract_text(path)


def test_pdf_parse():
    path = FIXTURES / "sample.pdf"
    if not path.exists():
        pytest.skip("sample.pdf fixture not available")
    text = extract_text(path)
    assert len(text) > 0


def test_docx_parse():
    path = FIXTURES / "sample.docx"
    if not path.exists():
        pytest.skip("sample.docx fixture not available")
    text = extract_text(path)
    assert len(text) > 0


def test_docx_table_extraction(tmp_path):
    from docx import Document

    path = tmp_path / "table.docx"
    doc = Document()
    doc.add_paragraph("工作经历")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "项目"
    table.rows[0].cells[1].text = "结果"
    table.rows[1].cells[0].text = "支付网关"
    table.rows[1].cells[1].text = "吞吐量提升 30%"
    doc.save(str(path))
    text = extract_text(path)
    assert "支付网关" in text
    assert "吞吐量提升 30%" in text


def test_normalize_text_collapses_blank_lines():
    assert normalize_text("  a \n\n  b \n\n\n c ") == "a\nb\nc"


def test_clean_resume_markdown_normalizes_bullets_and_invisible_chars():
    text = clean_resume_markdown("•\u200b 负责开发\u3000后端服务\n\u2022 使用 Python")
    assert "\u200b" not in text
    assert "\u3000" not in text
    assert "- 负责开发 后端服务" in text
    assert "- 使用 Python" in text


def test_clean_resume_markdown_merges_short_broken_lines_and_keeps_headings():
    text = clean_resume_markdown(
        "工作经历\n负责后端开发\n使用 Python\n项目经历\nRedis 缓存"
    )
    assert "工作经历\n负责后端开发 使用 Python" in text
    assert "项目经历\nRedis 缓存" in text


def test_clean_resume_markdown_preserves_english_lead_words():
    text = clean_resume_markdown(
        "Objective: backend engineer\nOverview of the project\n0 years experience"
    )
    assert "Objective: backend engineer" in text
    assert "Overview of the project" in text
    assert "0 years experience" in text


def test_structured_resume_sections():
    text = (
        "个人简介\n张三是名后端工程师\n\n"
        "# 工作经历\n负责高并发服务\n\n"
        "Skills\nPython\nFastAPI"
    )
    sections = structured_resume_sections(text)
    assert "个人简介" in sections
    assert "工作经历" in sections
    assert "Skills" in sections
    assert "Python" in sections["Skills"]


def test_unsupported_extension():
    with pytest.raises(FileParseError, match=r"\.xyz"):
        extract_text(FIXTURES / "dummy.xyz")


def test_file_not_found():
    with pytest.raises(FileParseError, match="not found"):
        extract_text(FIXTURES / "nonexistent.pdf")
