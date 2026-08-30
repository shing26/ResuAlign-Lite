from pathlib import Path

import pytest

FIXTURES = Path(__file__).parent / "fixtures"


@pytest.fixture(autouse=True)
def reset_shared_rate_limiters():
    """Keep per-host API rate limiters from starving later tests."""
    import resualign.api as api_module

    api_module._auth_rate_limiter.reset()
    api_module._analyze_rate_limiter.reset()
    api_module._import_rate_limiter.reset()
    yield


@pytest.fixture(autouse=True)
def stub_workbench_llm_probe(monkeypatch):
    """Never hit the network for the workbench pre-flight probe in tests.

    Phase A1 added a real one-token probe before queueing a workbench run.
    Tests build synthetic jobs/resumes and must not depend on connectivity
    or real credentials; the probe is stubbed as healthy by default, and
    individual tests that exercise A1's blocking branch re-stub it.
    """
    import resualign.api as api_module

    def _healthy_probe(tenant_id: str) -> tuple[bool, str]:
        return True, ""

    monkeypatch.setattr(
        api_module, "_probe_active_llm_quick", _healthy_probe
    )
    yield


class MockLLMClient:
    """Sequence-based fake LLM client for unit testing pipeline stages.

    Each call to ``chat_json`` consumes the next response from *responses*.
    Records calls for later assertion.
    """

    def __init__(self, responses=None):
        self.responses = list(responses or [
            {"score": 80, "skills": ["Python", "Java"], "issues": []},
        ])
        self.call_count = 0
        self.calls = []
        self.structured_calls = []

    def chat_json(self, system, user, model=None):
        self.calls.append((system[:50], user[:50]))
        if self.call_count < len(self.responses):
            r = self.responses[self.call_count]
            self.call_count += 1
            return r
        return {}

    def chat_structured(self, system, user, schema_model, model=None):
        self.structured_calls.append((system, user, schema_model, model))
        return self.chat_json(system, user, model=model)


class SchemaAwareLLMClient(MockLLMClient):
    """Fake LLM client that explicitly exposes chat_structured."""

    strict_provenance = True

    def chat_structured(self, system, user, schema_model, model=None):
        self.structured_calls.append((system, user, schema_model, model))
        return self.chat_json(system, user, model=model)


def _diag(score=80):
    return {"score": score, "skills": ["Python"], "issues": []}


def _profile():
    return {"must_have_skills": ["Java"], "nice_to_have_skills": ["Redis"],
            "soft_skills": [], "business_scenarios": ["Microservices"],
            "min_years_experience": None, "education_requirements": []}


def _gap():
    return {"missing_keywords": ["Redis"], "misaligned_emphasis": [],
            "strength_matches": ["Java"]}


def _jd_analysis():
    return {"jd_profile": _profile(), "gap_report": _gap()}



def _jd_profile_only():
    """Standalone JD profile (no gap_report wrapper)."""
    return _profile()


def _gap_only():
    """Standalone gap report (no jd_profile wrapper)."""
    return _gap()


def _tailor():
    return {"sections": {"experience": "Built services using Java"},
            "diffs": [{"type": "modify", "original": "Python dev",
                        "proposed": "Built services using Java",
                        "reason": "match", "confidence": "high",
                        "provenance": "Python dev",
                        "provenance_quote": "Python dev"}]}


def _eval():
    return {"jd_match_score": 82, "improvement": 12,
            "hallucination_detected": False, "hallucination_details": [],
            "gap_coverage": 0.75}


@pytest.fixture(scope="session")
def sample_pdf():
    """Dynamically generate a minimal PDF fixture using PyMuPDF."""
    path = FIXTURES / "sample.pdf"
    if not path.exists():
        import fitz
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "Python developer resume.\nJava experience.\nSpring Boot.")
        doc.save(str(path))
        doc.close()
    return path


@pytest.fixture(scope="session")
def sample_txt():
    return FIXTURES / "sample.txt"


@pytest.fixture(scope="session")
def sample_docx():
    """Dynamically generate a minimal DOCX fixture."""
    path = FIXTURES / "sample.docx"
    if not path.exists():
        from docx import Document
        doc = Document()
        doc.add_paragraph("Python developer resume.")
        doc.add_paragraph("Java experience.")
        doc.save(str(path))
    return path
